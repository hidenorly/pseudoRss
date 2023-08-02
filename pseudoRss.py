#!/usr/bin/env python3
# coding: utf-8

#   Copyright 2023 hidenorly
#
#   Licensed baseUrl the Apache License, Version 2.0 (the "License");
#   you may not use this file except in compliance with the License.
#   You may obtain a copy of the License at
#
#       http://www.apache.org/licenses/LICENSE-2.0
#
#   Unless required by applicable law or agreed to in writing, software
#   distributed baseUrl the License is distributed on an "AS IS" BASIS,
#   WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#   See the License for the specific language governing permissions and
#   limitations baseUrl the License.

import argparse
import os
import re
import string
import time
import datetime
import csv
import json
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from urllib.parse import urljoin
from urllib.parse import urlparse
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import StaleElementReferenceException
from selenium.common.exceptions import NoSuchElementException
try:
    import chromedriver_binary
except ImportError:
    pass


isVerbose = False

class HashCache:
    def __init__(self, cacheDir):
        self.cacheDir = cacheDir

    def getCacheFilename(self, url):
        parsed_url = urlparse(url)
        filename = parsed_url.netloc + parsed_url.path
        filename = re.sub(r'[^a-zA-Z0-9\-_.]', '_', filename)
        return os.path.join(self.cacheDir, filename)

    def ensureCacheStorage(self):
        if not os.path.exists(self.cacheDir):
            os.makedirs(self.cacheDir)

    def store(self, url, data):
        self.ensureCacheStorage()
        cachePath = self.getCacheFilename(url)

        dt_now = datetime.datetime.now()
        data["lastUpdate"] = dt_now.strftime("%Y-%m-%d %H:%M:%S")
        with open(cachePath, 'w', encoding='UTF-8') as f:
            json.dump(data, f, indent = 4, ensure_ascii=False)
            f.close()
        del data["lastUpdate"]

    def restore(self, url):
        result = {}

        cachePath = self.getCacheFilename(url)
        if os.path.exists( cachePath ):
            with open(cachePath, 'r', encoding='UTF-8') as f:
                result = json.load(f)
                if "lastUpdate" in result:
                    del result["lastUpdate"]

        return result


class WebLinkEnumerater:
    CONTROL_CHR_PATTERN = re.compile('[\x00-\x1f\x7f]')

    @staticmethod
    def isSameDomain(url1, url2, baseUrl=""):
        isSame = urlparse(url1).netloc == urlparse(url2).netloc
        isbaseUrl =  ( (baseUrl=="") or url2.startswith(baseUrl) )
        return isSame and isbaseUrl

    @staticmethod
    def getLinksByFactor(driver, pageUrl, byFactor=By.TAG_NAME, element='a', sameDomain=False, onlyTextExists=False):
        result = {}

        try:
            tag_name_elements = driver.find_elements(byFactor, element)
            for element in tag_name_elements:
                url = element.get_attribute('href')
                title = str(element.text).strip()
                title = WebLinkEnumerater.CONTROL_CHR_PATTERN.sub(' ', title)
                title = title.encode('utf-8', 'surrogatepass').decode('utf-8', 'ignore')
                if url:
                    if not sameDomain or WebLinkEnumerater.isSameDomain(pageUrl, url, pageUrl):
                        if not onlyTextExists or title:
                            result[url] = title
        except Exception as e: #except NoSuchElementException:
            pass

        return result

    @staticmethod
    def getLinks(driver, url, isSameDomain, onlyTextExists):
        result = {}

        if isVerbose:
            print(url)
        try:
            driver.get(url)
            result = WebLinkEnumerater.getLinksByFactor(driver, url, By.TAG_NAME, 'a', isSameDomain, onlyTextExists)
            result.update( WebLinkEnumerater.getLinksByFactor(driver, url, By.CSS_SELECTOR, 'a.post-link', isSameDomain, onlyTextExists) )
        except Exception as e: #except NoSuchElementException:
            print("Error at "+url)

        return result

    @staticmethod
    def getNewLinks(prevUrlList, newUrlList, stopIfExist = True):
        result = {}
        isFoundNewOne = False
        for aUrl, aTitle in newUrlList.items():
            found = aUrl in prevUrlList
            if not found:
                result[aUrl] = aTitle
                isFoundNewOne = True
            if isFoundNewOne and stopIfExist and found:
                break
        return result

class Reporter:
    def __init__(self, output = None):
        self.stream = None
        if output:
            self.stream = open(output, "a", encoding="utf-8")

    def _print(self, data):
        if self.stream:
            self.stream.write( str(data) + "\n" )
        else:
            print( str(data) )

    def print(self, data):
        for aUrl, aTitle in data["links"].items():
            self._print( str(aUrl) + ":" + str(aTitle) )

    def printHeader(self):
        pass

    def close(self):
        if self.stream:
            self.stream.close()
        self.stream = None

    def __del__(self):
        if self.stream:
            self.close()


class JsonReporter(Reporter):
    def __init__(self, output = None):
        super().__init__(output)
        self._print("[")

    def print(self, data):
        self._print( "\t" + json.dumps(data) + "," )

    def close(self):
        self._print("]")
        super().close()


class CsvReporter(Reporter):
    def __init__(self, output = None):
        super().__init__(output)

    def printHeader(self):
        self._print( "#site,url,title" )

    def print(self, data):
        for aUrl, aTitle in data["links"].items():
            self._print( str(data["site"]) + "," + str(aUrl) + "," + str(aTitle) )


class DocxReporter(Reporter):
    def __init__(self, output = "output.docx"):
        if output==None:
            output = "output.docx"
        self.document = Document()
        self.output = output
        if os.path.isfile(output):
            self.document = Document(output)

    def addTextWithLink(self, paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

    def printHeader(self):
        if self.document:
            doc = self.document

            # add date
            dt_now = datetime.datetime.now()
            today = dt_now.strftime("%d %B %Y")
            doc.add_heading(today, level=1)


    def print(self, data):
        if self.document:
            doc = self.document

            # add paragraph
            if data["links"] and "title" in data:
                doc.add_paragraph( data["title"] )

            # add list of links with title
            for aUrl, aTitle in data["links"].items():
                paragraph = doc.add_paragraph(style='List Bullet')
                self.addTextWithLink( paragraph, aTitle, aUrl )

    def close(self):
        if self.document:
            self.document.save(self.output)
        self.document = None

    def __del__(self):
        if self.document:
            self.close()



if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Pseudo RSS', formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument('pages', metavar='PAGE', type=str, nargs='*', help='Web URLs')
    parser.add_argument('-i', '--input', dest='input', type=str, default=None, help='list.csv url,title,sameDomain true or false,onlyTextExists true or false,newOnlyDiff true or false')
    parser.add_argument('-o', '--output', dest='output', type=str, default=None, help='Output filename')
    parser.add_argument('-c', '--cache', dest='cacheDir', type=str, default='~/.pseudoRss', help='Cache Dir')
    parser.add_argument('-s', '--sameDomain', dest='sameDomain', action='store_true', default=False, help='Specify if you want to restrict in the same url')
    parser.add_argument('-t', '--onlyTextExists', dest='onlyTextExists', action='store_true', default=False, help='Specify if you want to restrict text existing link')
    parser.add_argument('-d', '--diff', dest='diff', action='store_true', default=False, help='Specify if you want to list up new links')
    parser.add_argument('-n', '--newOnlyDiff', dest='newOnlyDiff', action='store_true', default=False, help='Specify if you want to enumerate new one only')
    parser.add_argument('-f', '--format', action='store', default="text", help='Set output format text or json or csv or docx')
    parser.add_argument('-v', '--verbose', dest='verbose', action='store_true', default=False, help='Specify if you want to enableverbose log output')
    args = parser.parse_args()
    isVerbose = args.verbose

    # setup selenium
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    driver = webdriver.Chrome(options=options)
    driver.set_window_size(1920, 1080)

    # cache, reporter
    cache = HashCache(os.path.expanduser(args.cacheDir))
    reporter = Reporter
    if args.format == "json":
        reporter = JsonReporter
    elif args.format == "csv":
        reporter = CsvReporter
    elif args.format == "docx":
        reporter = DocxReporter
    reporter = reporter(args.output)

    # input file
    pages = []
    if args.input and os.path.isfile(args.input):
        with open(args.input, 'r' ) as csvFile:
            data = csv.reader(csvFile)
            for rows in data:
                if args.verbose:
                    print("csv data is "+str(rows))
                if len(rows)>0:
                    aData = {
                        "url": rows[0],
                        "sameDomain": False,
                        "onlyTextExists": False,
                        "newOnlyDiff": False
                    }
                    if len(rows)>1 and rows[1]:
                        aData["title"] = rows[1]
                    if len(rows)>2 and rows[2] and rows[2].strip().lower()=="true":
                        aData["sameDomain"] = True
                    if len(rows)>3 and rows[3] and rows[3].strip().lower()=="true":
                        aData["onlyTextExists"] = True
                    if len(rows)>4 and rows[4] and rows[4].strip().lower()=="true":
                        aData["newOnlyDiff"] = True
                    pages.append(aData)
            csvFile.close()


    for aUrl in args.pages:
        pages.append({
                "url": aUrl,
                "sameDomain": args.sameDomain,
                "onlyTextExists": args.onlyTextExists,
                "newOnlyDiff": args.newOnlyDiff,
                "title" : ""
            })

    # enumeate link and output
    reporter.printHeader()

    for aPage in pages:
        aUrl = aPage["url"]
        if isVerbose:
            print("checking..."+aPage["title"]+" ("+aUrl+")...")
        urlList = WebLinkEnumerater.getLinks(driver, aUrl, aPage["sameDomain"], aPage["onlyTextExists"])
        listOut = urlList
        if args.diff:
            prevUrlList = cache.restore(aUrl)
            listOut = WebLinkEnumerater.getNewLinks(prevUrlList, urlList, aPage["newOnlyDiff"])
        cache.store(aUrl, urlList)

        outputData = {
            "site" : aPage["url"],
            "links": listOut
        }
        if "title" in aPage:
            outputData["title"] = aPage["title"]
        reporter.print( outputData )

    reporter.close()
    driver.quit()
